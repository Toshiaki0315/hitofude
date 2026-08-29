"""Word への書き出し（U-5。ユーザー要望 2026-08-29）。

**提出物が Word 指定**の場面は日本の実務で多く、PDF では代替できない。

**ざっくり作って手で整える前提**（PowerPoint への書き出しと同じ方針）。
凝った体裁は狙わず、見出し・段落・箇条書き・引用・表・コードが Word の
**標準スタイル**に乗っていればよい。受け取った人がテンプレートを当て
直せるほうが、こちらで飾り込むより役に立つ。

行の種類は `core/block_parser`、文中の装飾は `core/inline_scanner` が
決めている。**同じ解析を使う**——別に読み直すと、画面と書き出しで
解釈がずれる。
"""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from hitofude.core import frontmatter
from hitofude.core.block_parser import parse
from hitofude.core.inline_scanner import scan
from hitofude.core.models import BlockInfo, BlockType, InlineSpan, SpanType
from hitofude.core.paths import resolve_reference
from hitofude.core.table import split_cells

logger = logging.getLogger(__name__)

MONO_FAMILY = "Menlo"
"""コードの字。**受け取り側に無くても Word が代わりを当てる。**"""

CODE_POINT_SIZE = Pt(9.5)

_LIST_STYLES = {
    BlockType.BULLET_LIST_ITEM: "List Bullet",
    BlockType.ORDERED_LIST_ITEM: "List Number",
    BlockType.TASK_LIST_ITEM: "List Bullet",
}

MAX_IMAGE_WIDTH_IN = 6.0
"""絵の最大幅（インチ）。

A4 / Letter の既定の余白（左右 1 インチ）で本文に収まる幅。**大きい絵を
そのまま入れると紙からはみ出す**ので、ここで頭を打たせる。小さい絵は
そのままの大きさで入る。
"""

IMAGE_PLACEHOLDER = "［画像］"
"""説明（alt）の無い絵の代わり。

**黙って消さない。** 貼り付けた絵は `![](attachments/…)` の形で説明が
無いので、記号を外すと**何も残らない**——絵があったことすら伝わらない。
書き出しは絵を埋め込まない（PowerPoint と違い「手で整える前提」）ので、
置き場所だけ示す。
"""

CHECKED, UNCHECKED = "☑ ", "☐ "
"""チェックの印（レビュー指摘 2026-08-30）。

マーカーごと削っていたので、**済んだかどうかが消えて**ただの箇条書きに
なっていた。記号は `core/html` と同じ——書き出し先が違っても同じ印。
"""

# 文中で「地の文と違う見せ方」をするもの。ここに無い装飾は素の字で出す
_BOLD = {SpanType.STRONG, SpanType.STRONG_EM}
_ITALIC = {SpanType.EM, SpanType.STRONG_EM}
_STRIKE = {SpanType.STRIKE}

# **記号を落として中身だけ出すもの**（レビュー指摘 2026-08-30）。
# `[題](URL)` や `[[ノート]]` が生のまま出ていた。Word に持っていく人が
# 読むのは題名で、URL は本文の邪魔になる。
#
# **画像も同じ扱い**（レビュー指摘 2026-08-30 その 2）。入れていなかった
# ので、URL 側だけ落ちて `![図の名前]` という**壊れた断片**が出ていた。
# オートリンク（`<URL>`）も山括弧を外す
_UNWRAP = {SpanType.LINK_TEXT, SpanType.WIKI_LINK, SpanType.IMAGE, SpanType.AUTOLINK}

# 中身ごと落とすもの。`[題](URL)` の URL 側は題名の直後に続くので、
# 残すと `Qiita https://…` と 2 度出る
_DROP = {SpanType.LINK_URL}


def _body_of(line: str, info: BlockInfo) -> str:
    """行頭マーカーを外した中身。`## 題` なら `題`。"""
    return line[info.marker_len :] if info.marker_len else line


def _runs(text: str) -> list[tuple[str, SpanType | None]]:
    """文中を `(字, 装飾)` に割る。**記号は落とす**（`**太字**` → 太字）。

    重なった装飾は外側だけを見る。Word の run は入れ子にできないので、
    どのみち 1 段しか表せない。
    """
    wanted = _BOLD | _ITALIC | _STRIKE | _UNWRAP | _DROP | {SpanType.CODE}
    spans: list[InlineSpan] = [span for span in scan(text) if span.type in wanted]
    found: list[tuple[str, SpanType | None]] = []
    at = 0
    for span in sorted(spans, key=lambda item: item.open_start):
        if span.open_start < at:
            continue  # 入れ子。外側だけを見る
        if span.open_start > at:
            found.append((text[at : span.open_start], None))
        if span.type not in _DROP:
            body = text[span.open_end : span.close_start]
            if span.type is SpanType.IMAGE and not body:
                body = IMAGE_PLACEHOLDER
            found.append((body, span.type))
        at = span.close_end
    if at < len(text):
        found.append((text[at:], None))
    return [(body, kind) for body, kind in found if body]


def _picture_size(path: Path) -> "Inches | None":
    """入れる幅。紙からはみ出さないところで頭を打たせる。

    元の大きさが分からない（読めない）ときは `None` を返し、
    python-docx に元の大きさで置いてもらう。
    """
    from PySide6.QtGui import QImageReader

    size = QImageReader(str(path)).size()
    if not size.isValid() or size.width() <= 0:
        return None
    # 96dpi と見なす（Word の既定）。それより大きければ幅を頭打ちにする
    natural = size.width() / 96
    return Inches(min(natural, MAX_IMAGE_WIDTH_IN))


def _add_picture(paragraph, url: str, base_path: Path | None) -> bool:
    """絵を段落へ入れる。入れられたら True。

    **書き出しを止めない**（PowerPoint への書き出しと同じ方針）。
    保管フォルダの外・見つからない・読めない、のどれでも飛ばすだけ。
    """
    resolved = resolve_reference(base_path, url)
    if resolved is None or not resolved.is_file():
        return False
    try:
        paragraph.add_run().add_picture(str(resolved), width=_picture_size(resolved))
    except Exception as error:  # python-docx は形式ごとに別の例外を投げる
        logger.warning("画像を入れられなかった: %s (%s)", resolved, error)
        return False
    return True


def _write_runs(paragraph, text: str, base_path: Path | None = None) -> None:
    for body, kind in _runs(text):
        # **絵そのものを入れる**（ユーザー要望 2026-08-30）。入らなかった
        # ときだけ、説明か在処の印を字で残す
        if kind is SpanType.IMAGE and _add_picture(paragraph, _image_url(text, body), base_path):
            continue
        run = paragraph.add_run(body)
        if kind in _BOLD:
            run.bold = True
        if kind in _ITALIC:
            run.italic = True
        if kind in _STRIKE:
            run.font.strike = True
        if kind is SpanType.CODE:
            run.font.name = MONO_FAMILY


def _image_url(text: str, alt: str) -> str:
    """`![alt](URL)` の URL。走査は説明までしか返さないので、後ろを読む。"""
    marker = f"![{alt}](" if alt != IMAGE_PLACEHOLDER else "![]("
    at = text.find(marker)
    if at < 0:
        return ""
    end = text.find(")", at + len(marker))
    return text[at + len(marker) : end] if end > 0 else ""


def _cells(line: str) -> list[str]:
    """表の 1 行をセルに割る。

    **両端の空セルを落とす。** `split_cells` は外側の `|` のぶんも返すので
    （`['', ' 項目 ', ' 担当 ', '']`）、そのままだと空の列が 2 つ増える。
    **中の空セルは残す**——`| あ | | う |` の真ん中は書き手が空けたもの。
    """
    found = [cell.strip() for cell in split_cells(line)]
    if found and not found[0]:
        found = found[1:]
    if found and not found[-1]:
        found = found[:-1]
    return found


def _write_table(document, rows: list[list[str]]) -> None:
    """表を **Word の表**として出す。段落にしない。"""
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for y, row in enumerate(rows):
        for x in range(width):
            cell = table.cell(y, x)
            cell.text = ""
            _write_runs(cell.paragraphs[0], row[x] if x < len(row) else "")


def _write_code(document, lines: list[str]) -> None:
    """コードは**1 段落にまとめる**。1 行ずつ足すと段落の間隔で間延びする。"""
    if not lines:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = MONO_FAMILY
    run.font.size = CODE_POINT_SIZE


def write_docx(path: Path, text: str, *, base_path: Path | None = None) -> Path:
    """本文を `.docx` として書く。書いた先を返す。

    **front matter は出さない**（作成日時と id は読む人に要らない）。

    `base_path` は絵を探す起点（保管フォルダ）。渡さなければ絵は入らず、
    説明だけが字で残る。**外は読まない**（`resolve_reference` の約束）。
    """
    document = Document()
    body = frontmatter.split(text).body
    lines = body.split("\n")
    blocks = parse(body)

    code: list[str] = []
    table: list[list[str]] = []
    for info, line in zip(blocks, lines, strict=False):
        # 表とコードは**続きを溜めてから**まとめて出す
        # **区切り行（`|---|`）で切らない。** 表の一部なので、ここで
        # 吐き出すとヘッダだけの表と本体の表に割れる
        if info.type not in _TABLE_TYPES and table:
            _write_table(document, table)
            table = []
        if info.type not in _CODE_TYPES and code:
            _write_code(document, code)
            code = []

        match info.type:
            case BlockType.HEADING:
                document.add_heading(_body_of(line, info).strip(), level=info.level)
            case BlockType.CODE_FENCE_BODY:
                code.append(line)
            case BlockType.CODE_FENCE_OPEN | BlockType.CODE_FENCE_CLOSE:
                pass  # 記号は出さない
            case BlockType.TABLE_ROW:
                table.append(_cells(line))
            case BlockType.TABLE_DELIMITER:
                pass  # `|---|` は罫線の指定。中身ではない
            case BlockType.BLOCKQUOTE:
                _write_runs(
                    document.add_paragraph(style="Quote"),
                    _body_of(line, info).strip(),
                    base_path,
                )
            case kind if kind in _LIST_STYLES:
                body = _body_of(line, info).strip()
                if kind is BlockType.TASK_LIST_ITEM:
                    # **済んだかどうかを落とさない。** マーカーは削られて
                    # いるので、印を先頭に置き直す
                    body = (CHECKED if info.checked else UNCHECKED) + body
                _write_runs(document.add_paragraph(style=_LIST_STYLES[kind]), body, base_path)
            case BlockType.HORIZONTAL_RULE:
                document.add_paragraph()
            case BlockType.PARAGRAPH:
                _write_runs(document.add_paragraph(), line.strip(), base_path)
            case _:
                pass  # 空行・front matter・数式の記号など

    _write_table(document, table)
    _write_code(document, code)
    document.save(str(path))
    return path


_TABLE_TYPES = frozenset({BlockType.TABLE_ROW, BlockType.TABLE_DELIMITER})

_CODE_TYPES = frozenset(
    {BlockType.CODE_FENCE_BODY, BlockType.CODE_FENCE_OPEN, BlockType.CODE_FENCE_CLOSE}
)
