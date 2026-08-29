"""Word 書き出しの入口（U-5）。"""

import pytest

pytestmark = pytest.mark.gui


class TestEntry:
    def test_書き出しメニューにある(self, window) -> None:
        from hitofude.ui.commands import commands

        labels = [c.label for c in commands(window.menuBar())]
        assert "Word…" in labels

    def test_書き出せる(self, window, tmp_path, monkeypatch) -> None:
        from docx import Document

        note = window._vault.create("会議メモ", "# 会議メモ\n\n決めたこと\n")
        window._db.upsert_note(note, window._vault.root)
        window.refresh()
        window.open_and_select(note.path)

        target = tmp_path / "会議メモ.docx"
        monkeypatch.setattr(
            "hitofude.ui.export_actions.QFileDialog.getSaveFileName",
            lambda *a, **k: (str(target), ""),
        )
        assert window.export_docx() == target
        assert any("決めたこと" in p.text for p in Document(str(target)).paragraphs)


class TestImagesFromVault:
    """書き出しの入口から保管フォルダを渡す（ユーザー要望 2026-08-30）。

    渡さないと絵が入らない。**画面と同じ起点**（`_vault.root`）で探す。
    """

    def test_絵が入る(self, window, tmp_path, monkeypatch) -> None:
        from docx import Document
        from tests.editor.test_docx_export import make_png

        make_png(window._vault.root / "attachments" / "zu.png", width=30, height=20, color=0x224466)

        note = window._vault.create("絵入り", "# 絵入り\n\n![図](attachments/zu.png)\n")
        window._db.upsert_note(note, window._vault.root)
        window.refresh()
        window.open_and_select(note.path)

        target = tmp_path / "絵入り.docx"
        monkeypatch.setattr(
            "hitofude.ui.export_actions.QFileDialog.getSaveFileName",
            lambda *a, **k: (str(target), ""),
        )
        assert window.export_docx() == target
        document = Document(str(target))
        drawings = sum(
            len(
                p._p.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
                )
            )
            for p in document.paragraphs
        )
        assert drawings == 1, "絵が入っていない"
