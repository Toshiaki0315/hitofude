"""Word への書き出し（U-5。ユーザー要望 2026-08-29）。

**提出物が Word 指定**の場面は日本の実務で多く、PDF では代替できない。
Bear・Typora・iA Writer にはある。

**ざっくり作って手で整える前提**（PowerPoint への書き出しと同じ方針）。
凝った体裁は狙わず、見出し・段落・箇条書き・表・コードが Word の
スタイルに乗っていればよい。

書き出したものを **python-docx で読み返して**確かめる。作れただけでは
中身の保証にならない。
"""

from pathlib import Path

import pytest
from docx import Document

from hitofude.editor.docx_export import write_docx

pytestmark = []

NOTE = """# 会議メモ

決めたことを **太字** と *斜体* と `コード` で書く。

## 次にやること

- 見積もりを出す
- 日程を決める

1. まず確認
2. つぎに連絡

> 引用した一言

```python
print(1)
```

| 項目 | 担当 |
| --- | --- |
| 見積もり | 山田 |
| 日程 | 佐藤 |
"""


@pytest.fixture
def written(tmp_path):
    target = tmp_path / "会議メモ.docx"
    write_docx(target, NOTE)
    return Document(str(target))


def make_png(path: Path, *, width: int = 40, height: int = 30, color: int = 0x336699) -> Path:
    """試験用の絵を 1 枚置く。

    **書き方を 1 か所にまとめる**（レビュー指摘 2026-08-30）。同じ数行が
    6 箇所に写っていた。大きさは埋め込みの寸法を見る試験が指定する。
    """
    from PySide6.QtGui import QImage

    path.parent.mkdir(parents=True, exist_ok=True)
    picture = QImage(width, height, QImage.Format.Format_RGB32)
    picture.fill(color)
    assert picture.save(str(path)), f"絵を作れなかった: {path}"
    return path


def texts(document) -> list[str]:
    return [p.text for p in document.paragraphs]


class TestStructure:
    def test_見出しがスタイルに乗る(self, written) -> None:
        found = next(p for p in written.paragraphs if p.text == "会議メモ")
        assert found.style.name.startswith("Heading")

    def test_見出しの深さが伝わる(self, written) -> None:
        top = next(p for p in written.paragraphs if p.text == "会議メモ")
        second = next(p for p in written.paragraphs if p.text == "次にやること")
        assert top.style.name != second.style.name

    def test_段落が出る(self, written) -> None:
        assert any("決めたこと" in text for text in texts(written))

    def test_箇条書きがスタイルに乗る(self, written) -> None:
        found = next(p for p in written.paragraphs if p.text == "見積もりを出す")
        assert "List" in found.style.name

    def test_番号付きも出る(self, written) -> None:
        assert any("まず確認" in text for text in texts(written))

    def test_引用がスタイルに乗る(self, written) -> None:
        found = next(p for p in written.paragraphs if "引用した一言" in p.text)
        assert found.style.name in {"Quote", "Intense Quote"}

    def test_表が表になる(self, written) -> None:
        """**表を段落にしない。** Word の表として出す。"""
        assert written.tables
        table = written.tables[0]
        assert table.cell(0, 0).text == "項目"
        assert table.cell(1, 1).text == "山田"

    def test_コードは等幅(self, written) -> None:
        found = next(p for p in written.paragraphs if "print(1)" in p.text)
        assert found.runs and found.runs[0].font.name


class TestInline:
    """文中の装飾。**記号は残さない**（`**太字**` ではなく太字）。"""

    def test_太字になる(self, written) -> None:
        paragraph = next(p for p in written.paragraphs if "決めたこと" in p.text)
        assert any(run.bold and run.text == "太字" for run in paragraph.runs)

    def test_斜体になる(self, written) -> None:
        paragraph = next(p for p in written.paragraphs if "決めたこと" in p.text)
        assert any(run.italic and run.text == "斜体" for run in paragraph.runs)

    def test_記号が残らない(self, written) -> None:
        paragraph = next(p for p in written.paragraphs if "決めたこと" in p.text)
        assert "**" not in paragraph.text
        assert "`" not in paragraph.text


class TestQuiet:
    def test_front_matterは出さない(self, tmp_path) -> None:
        target = tmp_path / "a.docx"
        write_docx(target, "---\ncreated: 2026-08-29\n---\n\n# 題\n\n本文\n")
        assert not any("created" in text for text in texts(Document(str(target))))

    def test_空でも書ける(self, tmp_path) -> None:
        target = tmp_path / "b.docx"
        write_docx(target, "")
        assert target.is_file()


class TestTaskList:
    """チェックの状態を落とさない（レビュー指摘 2026-08-30）。

    `- [x] 済み` のマーカーごと削っていたので、**済んだかどうかが
    消えて**ただの箇条書きになっていた。記号は `core/html` と揃える
    （書き出し先が違っても同じ印）。
    """

    @pytest.fixture
    def tasks(self, tmp_path):
        target = tmp_path / "tasks.docx"
        write_docx(target, "# やること\n\n- [ ] まだ\n- [x] 済み\n")
        return Document(str(target))

    def test_未了の印が出る(self, tasks) -> None:
        assert any("☐" in text and "まだ" in text for text in texts(tasks))

    def test_済みの印が出る(self, tasks) -> None:
        assert any("☑" in text and "済み" in text for text in texts(tasks))

    def test_ふつうの箇条書きには付かない(self, tmp_path) -> None:
        target = tmp_path / "plain.docx"
        write_docx(target, "# 一覧\n\n- ただの項目\n")
        assert not any("☐" in text for text in texts(Document(str(target))))


class TestInlineMarkup:
    """記号を生のまま出さない（レビュー指摘 2026-08-30）。"""

    @pytest.fixture
    def written(self, tmp_path):
        target = tmp_path / "inline.docx"
        write_docx(
            target,
            "# 見本\n\n[Qiita](https://qiita.com) と [[会議メモ]] と ~~取り消し~~ を書く。\n",
        )
        return Document(str(target))

    def test_リンクは題名だけ出す(self, written) -> None:
        paragraph = next(p for p in written.paragraphs if "Qiita" in p.text)
        assert "https://qiita.com" not in paragraph.text
        assert "](" not in paragraph.text

    def test_ノートのリンクも記号を外す(self, written) -> None:
        paragraph = next(p for p in written.paragraphs if "会議メモ" in p.text)
        assert "[[" not in paragraph.text

    def test_打ち消し線になる(self, written) -> None:
        paragraph = next(p for p in written.paragraphs if "取り消し" in p.text)
        assert any(run.font.strike and run.text == "取り消し" for run in paragraph.runs)
        assert "~~" not in paragraph.text


class TestImageAndAutolink:
    """画像とオートリンクの記号（レビュー指摘 2026-08-30）。"""

    @pytest.fixture
    def written(self, tmp_path):
        target = tmp_path / "media.docx"
        write_docx(
            target,
            "# 見本\n\n"
            "図は ![図の名前](zu.png) です。\n\n"
            "貼った絵は ![](attachments/a.png) です。\n\n"
            "参照は <https://example.com> です。\n",
        )
        return Document(str(target))

    def test_画像は説明だけ残る(self, written) -> None:
        """**壊れた断片を出さない。** `![図の名前]` のまま出ていた（実測）。"""
        paragraph = next(p for p in written.paragraphs if "図の名前" in p.text)
        assert "![" not in paragraph.text
        assert "]" not in paragraph.text
        assert "zu.png" not in paragraph.text

    def test_説明の無い絵は在処を示す(self, written) -> None:
        """**黙って消さない。** 貼った絵は `![](…)` で説明が無い。
        何も出さないと、絵があったことすら伝わらない。
        """
        from hitofude.editor.docx_export import IMAGE_PLACEHOLDER

        assert any(IMAGE_PLACEHOLDER in p.text for p in written.paragraphs)

    def test_オートリンクは山括弧を外す(self, written) -> None:
        paragraph = next(p for p in written.paragraphs if "example.com" in p.text)
        assert "<" not in paragraph.text
        assert ">" not in paragraph.text
        assert "https://example.com" in paragraph.text


class TestEmbeddedImages:
    """絵そのものを Word へ入れる（ユーザー要望 2026-08-30）。

    提出物として使うなら、絵が抜けていては用をなさない。PowerPoint への
    書き出しは既に入れているので、**同じ作法**（`resolve_reference` で
    保管フォルダの中だけを解決し、読めなければ飛ばす）に揃える。
    """

    @pytest.fixture
    def vault(self, tmp_path):
        root = tmp_path / "v"
        make_png(root / "attachments" / "zu.png")
        return root

    def written(self, tmp_path, vault, text: str):
        target = tmp_path / "out.docx"
        write_docx(target, text, base_path=vault)
        return Document(str(target))

    def pictures(self, document) -> int:
        """入っている絵の数。段落の中の `<w:drawing>` を数える。"""
        return sum(
            len(
                paragraph._p.findall(
                    ".//" + "{http://schemas.openxmlformats.org/"
                    "drawingml/2006/wordprocessingDrawing}anchor"
                )
            )
            + len(
                paragraph._p.findall(
                    ".//" + "{http://schemas.openxmlformats.org/"
                    "drawingml/2006/wordprocessingDrawing}inline"
                )
            )
            for paragraph in document.paragraphs
        )

    def test_絵が入る(self, tmp_path, vault) -> None:
        """**これが本題。**"""
        document = self.written(tmp_path, vault, "# 見本\n\n![図](attachments/zu.png)\n")
        assert self.pictures(document) == 1

    def test_文の途中の絵も入る(self, tmp_path, vault) -> None:
        document = self.written(tmp_path, vault, "# 見本\n\n前 ![図](attachments/zu.png) 後\n")
        assert self.pictures(document) == 1
        assert any("前" in p.text and "後" in p.text for p in document.paragraphs)

    def test_無い絵は在処を示すだけ(self, tmp_path, vault) -> None:
        """**書き出しを止めない**（PowerPoint と同じ方針）。"""
        from hitofude.editor.docx_export import IMAGE_PLACEHOLDER

        document = self.written(tmp_path, vault, "# 見本\n\n![](attachments/無い.png)\n")
        assert self.pictures(document) == 0
        assert any(IMAGE_PLACEHOLDER in p.text for p in document.paragraphs)

    def test_保管フォルダの外は入れない(self, tmp_path, vault) -> None:
        """**外は読まない**（`resolve_reference` の約束）。"""
        document = self.written(tmp_path, vault, "# 見本\n\n![外](../外.png)\n")
        assert self.pictures(document) == 0

    def test_置き場を渡さなければ今までどおり(self, tmp_path, vault) -> None:
        """`base_path` 無しでも落ちない（説明だけ残る）。"""
        target = tmp_path / "none.docx"
        write_docx(target, "# 見本\n\n![図](attachments/zu.png)\n")
        assert any("図" in p.text for p in Document(str(target)).paragraphs)


class TestManyImages:
    """同じ行に絵が 2 つあるとき（レビュー指摘 2026-08-30）。

    URL を**行の中から文字列で探して**いたので、同じ説明の絵が並ぶと
    2 枚目も 1 枚目の URL を引いていた（実測: `a.png` が 2 回）。
    走査（`scan`）が既に正しい URL を持っている。
    """

    @pytest.fixture
    def vault(self, tmp_path):
        root = tmp_path / "v"
        for name, color in (("a.png", 0x882222), ("b.png", 0x228822)):
            make_png(root / "attachments" / name, width=20, height=15, color=color)
        return root

    def test_それぞれの絵が入る(self, tmp_path, vault) -> None:
        """**これが本題。** 同じ絵を 2 度入れない。"""
        from hitofude.editor.docx_export import _runs

        text = "![図](attachments/a.png) と ![図](attachments/b.png)"
        urls = [payload for _body, kind, payload in _runs(text) if kind is not None and payload]
        assert urls == ["attachments/a.png", "attachments/b.png"]

    def test_書き出しにも2枚入る(self, tmp_path, vault) -> None:
        target = tmp_path / "two.docx"
        write_docx(
            target,
            "# 見本\n\n![図](attachments/a.png) と ![図](attachments/b.png)\n",
            base_path=vault,
        )
        document = Document(str(target))
        ns = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        assert sum(len(p._p.findall(".//" + ns)) for p in document.paragraphs) == 2


class TestImagesInTable:
    """表のセルの絵も入れる（レビュー指摘 2026-08-30）。

    `_write_table` へ置き場を渡していなかったので、セルの中だけ
    `［画像］` のままだった。
    """

    def test_セルの絵が入る(self, tmp_path) -> None:
        root = tmp_path / "v"
        make_png(root / "attachments" / "icon.png", width=20, height=15, color=0x224488)

        target = tmp_path / "table.docx"
        write_docx(
            target,
            "# 見本\n\n| 印 | 名前 |\n| --- | --- |\n| ![印](attachments/icon.png) | あ |\n",
            base_path=root,
        )
        table = Document(str(target)).tables[0]
        ns = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        cell = table.cell(1, 0)
        assert sum(len(p._p.findall(".//" + ns)) for p in cell.paragraphs) == 1


class TestTallImage:
    """縦に長い絵を紙からはみ出させない（レビュー指摘 2026-08-30）。

    幅にしか上限が無かったので、細長い画面写真（スマホの長尺）が
    紙の高さを超えていた。
    """

    def test_高さも収める(self, tmp_path) -> None:
        from docx.shared import Inches

        from hitofude.editor.docx_export import MAX_IMAGE_HEIGHT_IN, _picture_size

        # 1:10 の細長い絵
        path = make_png(tmp_path / "v" / "tall.png", width=300, height=3000, color=0x333333)

        width = _picture_size(path)
        assert width is not None
        # 幅 : 高さ = 1 : 10 なので、高さの上限から幅が決まる
        assert width <= Inches(MAX_IMAGE_HEIGHT_IN / 10) + Inches(0.01)

    def test_ふつうの絵はそのまま(self, tmp_path) -> None:
        from docx.shared import Inches

        from hitofude.editor.docx_export import _picture_size

        # 96dpi で 2 インチぶん
        path = make_png(tmp_path / "v" / "normal.png", width=192, height=96, color=0x777777)
        assert _picture_size(path) == Inches(2.0)
